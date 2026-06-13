import os
import re
import json
import openpyxl
from docx import Document
from aiogram import Router, F, types
from aiogram.filters import Command
from aiogram.fsm.context import FSMContext
from aiogram.utils.keyboard import ReplyKeyboardBuilder
import config.database as db
from filters.admin_filter import IsAdmin
from states.admin_states import AdminStates

router = Router()

def get_admin_keyboard():
    builder = ReplyKeyboardBuilder()
    builder.button(text="📥 Istalgan formatda fayl yuklash")
    builder.button(text="📊 Statistika")
    builder.adjust(1, 1)
    return builder.as_markup(resize_keyboard=True)

@router.message(Command("start"), IsAdmin())
async def admin_start(message: types.Message):
    await message.answer(f"👨‍💻 Xush kelibsiz Admin, {message.from_user.full_name}!", reply_markup=get_admin_keyboard())

@router.message(F.text == "📊 Statistika", IsAdmin())
async def show_stats(message: types.Message):
    async with db.db_conn.execute("SELECT COUNT(*) FROM students") as cursor:
        total = (await cursor.fetchone())[0]
    await message.answer(f"📊 **Tizim statistikasi:**\n\nJami bazadagi unikal talabalar: {total} ta", parse_mode="Markdown")

@router.message(F.text == "📥 Istalgan formatda fayl yuklash", IsAdmin())
async def universal_file_start(message: types.Message, state: FSMContext):
    await state.set_state(AdminStates.waiting_for_universal)
    await message.answer(
        "📁 Menga istalgan formatdagi faylni yuboring (`.txt`, `.csv`, `.xlsx`, `.docx`, `.json`)\n\n"
        "✨ **Dublikatga qarshi tizim yoqilgan:** Agar talaba bazada allaqachon mavjud bo'lsa, u qayta qo'shilmaydi va faqat 1 marta chiqadi!",
        reply_markup=types.ReplyKeyboardRemove()
    )

def clean_text_to_dict(line_str: str):
    """ Qatordagi matndan unikal ma'lumotlarni saralab olish """
    if not line_str or len(line_str.strip()) < 5:
        return None
        
    # 1. Telefon qidirish
    phone_match = re.search(r'(?:\+?998)?\s?\(?\d{2}\)?\s?\d{3}\s?\d{2}\s?\d{2}', line_str)
    phone = phone_match.group(0).replace(" ", "") if phone_match else "Mavjud emas"
    if phone_match:
        line_str = line_str.replace(phone_match.group(0), "")
        
    # 2. Kurs qidirish
    course_match = re.search(r'([1-4])\s?(-?kurs)?', line_str, re.IGNORECASE)
    course = int(course_match.group(1)) if course_match else 1
    if course_match:
        line_str = line_str.replace(course_match.group(0), "")
        
    # 3. F.I.SH qidirish
    words = re.findall(r'[A-Z\'`‘a-zА-ЯЁа-яё]+', line_str)
    if len(words) < 2:
        return None
        
    full_name = " ".join(words[:3]).strip()
    remains = " ".join(words[3:]).strip() if len(words) > 3 else ""
    
    return {
        "full_name": full_name,
        "course": course,
        "phone_number": phone,
        "extra": f"Qo'shimcha ma'lumotlar: {remains}" if remains else "Fayldan yozib olindi"
    }

@router.message(AdminStates.waiting_for_universal, F.document)
async def process_any_file(message: types.Message, state: FSMContext, bot):
    file_name = message.document.file_name.lower()
    file_ext = os.path.splitext(file_name)[1]
    
    file = await bot.get_file(message.document.file_id)
    destination = f"temp_file{file_ext}"
    await bot.download_file(file.file_path, destination)
    
    await message.answer("🔄 Fayl tahlil qilinmoqda, dublikatlar tozalanmoqda...")
    students_list = []
    seen_keys = set() # Fayl ichidagi bir xil ismlarni o'sha zahoti o'chirish uchun
    
    try:
        # 1. EXCEL FORMAT (.xlsx)
        if file_ext == '.xlsx':
            wb = openpyxl.load_workbook(destination, data_only=True)
            sheet = wb.active
            for row in sheet.iter_rows(values_only=True):
                row_str = " ".join([str(cell) for cell in row if cell is not None])
                res = clean_text_to_dict(row_str)
                if res:
                    res["extra"] = f"Excel qatori: {row_str}"
                    # Unikallik kaliti: Ism + Telefon
                    key = (res["full_name"].lower(), res["phone_number"])
                    if key not in seen_keys:
                        seen_keys.add(key)
                        students_list.append(res)
                    
        # 2. WORD FORMAT (.docx)
        elif file_ext == '.docx':
            doc = Document(destination)
            # Jadvallardan o'qish
            for table in doc.tables:
                for row in table.rows:
                    row_str = " ".join([cell.text.strip() for cell in row.cells])
                    res = clean_text_to_dict(row_str)
                    if res:
                        key = (res["full_name"].lower(), res["phone_number"])
                        if key not in seen_keys:
                            seen_keys.add(key)
                            students_list.append(res)
            # Matnlardan o'qish
            for para in doc.paragraphs:
                res = clean_text_to_dict(para.text)
                if res:
                    key = (res["full_name"].lower(), res["phone_number"])
                    if key not in seen_keys:
                        seen_keys.add(key)
                        students_list.append(res)
                
        # 3. TEXT YOKI CSV FORMAT (.txt, .csv)
        elif file_ext in ['.txt', '.csv']:
            with open(destination, "r", encoding="utf-8", errors="ignore") as f:
                for line in f:
                    res = clean_text_to_dict(line)
                    if res:
                        key = (res["full_name"].lower(), res["phone_number"])
                        if key not in seen_keys:
                            seen_keys.add(key)
                            students_list.append(res)
                    
        # 4. JSON FORMAT (.json)
        elif file_ext == '.json':
            with open(destination, "r", encoding="utf-8") as f:
                data = json.load(f)
                if isinstance(data, list):
                    for item in data:
                        res = clean_text_to_dict(str(item))
                        if res:
                            key = (res["full_name"].lower(), res["phone_number"])
                            if key not in seen_keys:
                                seen_keys.add(key)
                                students_list.append(res)

        # 🚀 BAZAGA FAQAT YANGI (MAVJUD BO'LMAGAN) TALABALARNI YOZISH 🚀
        inserted_count = 0
        ignored_count = 0
        
        for s in students_list:
            # Avval bazada bu ismli talaba bor-yo'qligini tekshiramiz
            async with db.db_conn.execute(
                "SELECT id FROM students WHERE LOWER(full_name) = ? AND phone_number = ?", 
                (s['full_name'].lower(), s['phone_number'])
            ) as check_cursor:
                already_exists = await check_cursor.fetchone()
                
            if not already_exists:
                # Agar bazada bo'lmasa, yangi talaba sifatida qo'shamiz
                await db.db_conn.execute('''
                    INSERT INTO students (full_name, course, phone_number, extra_info)
                    VALUES (?, ?, ?, ?)
                ''', (s['full_name'], s['course'], s['phone_number'], s['extra']))
                inserted_count += 1
            else:
                ignored_count += 1
            
        await db.db_conn.commit()
        await message.answer(
            f"✅ Fayl muvaffaqiyatli yuklandi!\n\n"
            f"📥 Yangi qo'shildi: **{inserted_count}** ta talaba.\n"
            f"⚠️ Rad etildi (bazada borligi uchun): **{ignored_count}** ta.", 
            reply_markup=get_admin_keyboard(), 
            parse_mode="Markdown"
        )
        
    except Exception as e:
        await message.answer(f"❌ Faylni qayta ishlashda kutilmagan xatolik: {e}", reply_markup=get_admin_keyboard())
    finally:
        if os.path.exists(destination): 
            os.remove(destination)
            
    await state.clear()